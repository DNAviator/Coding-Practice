from PyPDF2 import PdfReader
from docx import Document


def format_page(page_text, index):
    # invalid case
    if index <= 0:
        return ""

    # Names of the specfic drivers or practices
    names = {
        1: "Aspires to Greatness",
        2: "Learns New Jobs and Tasks Quickly",
        3: "Always Embraces and Leads Change",
        4: "Seeks New Experiences and is an Exemplar for Being a Life Long Learner",
        5: "Seeks Variety and Diversity of Experience",
        6: "Digests Systems and Complexity Comfortably",
        7: "Constantly Builds Self-Awareness and Self-Regulation",
        8: "Networks Effectively in All Directions to Get Great Things Done",
        9: "Motivated by Challenges, Adversity, Ambiguity, and Intensity",
        10: "Able to See Past Mountains and Around Corners to Anticipate Events and Trends and Identify Potential Obstacles",
        11: "Takes the Initiative; Looks for Creative and Innovative Approaches",
        12: "Views Everything Broadly; Has Optimistic Expectations and Assumes a Growth Mindset",
        13: "Self-Awareness and Development",
        14: "Strong Aspirations",
        15: "Uncertainty and Ambiguity Management",
        16: "Adjustive, Adaptive, and Resilient",
        17: "Emotion and Stress Management",
        18: "Takes Risks and the Initiative",
        19: "Mindfulness Management",
        20: "Self-Confidence",
        21: "Challenge Seeking",
        22: "Reading People, Groups, and Enterprises",
        23: "Managing Conflict Productively",
        24: "Motivating and Influencing",
        25: "Managing People Differently",
        26: "Networking and Collaborating",
        27: "Cognitive Processing Capacity and Speed",
        28: "Fluid Intelligence",
        29: "Cognitive Complexity",
        30: "Systems Understanding",
        31: "Essence Detector",
        32: "Growth Mindset",
        33: "Global and Broad Vision and Perspective",
        34: "Fostering Creativity and Innovation",
        35: "Quick Learner",
        36: "Achievement Driven",
        37: "Beginner’s Mindset"
    }

    # General patterns of replacement for translate
    replacements = {
        "KSAL": "KSA - Leaders",
        "KSAM": "KSA - Managers",
        "KSAI": "KSA - Individual Contributors",
        "KSAT": "KSA - Teams",
        "KSAP": "KSA - Potentials",
        "Definition:\n": f"Definition of {names[index]}:\n",
        "Observables and Proofs:\n": f"Observables and Proofs of {names[index]}:\n",
        f"Some indicators that Driver/Marker {index}:": "Some indicators that the",
        f" D&M {index},": "",
        f" D&M {index}:": ""
    }

    # for this page nothing of importance exists and all of it needs to be replaced
    if "\nC\n" in page_text and index <= 12:
        return f"Degree of difficulty of {names[index]} Driver/Marker:"

    # replaces all selected items in dictionary
    for old, new in replacements.items():
        page_text = page_text.replace(old, new)

    lineArray = page_text.splitlines()
    if "D/M" in lineArray[-2]:
        lineArray[-1] = ""
        lineArray[-2] = lineArray[-2][0:-3]

    return "\n".join(lineArray[2:])

def extract_text_from_pdf(pdf_file_path, page_groups):
    """Extracts text from the specified PDF file, ignoring images."""

    pdf_reader = PdfReader(pdf_file_path)
    document = Document()
    index = 0
    for page_num in range(len(pdf_reader.pages)):

        # Adds page to document as a paragraph
        page = pdf_reader.pages[page_num]

        # extracts text and formats it before adding to document
        document.add_paragraph(format_page(page.extract_text(), index))

        # When an endpoint is reached file is saved a new file is opened
        if page_num+2 in page_groups:

            # labels file with correct numbering
            if page_num > 150:
                # practices portion of the pdf
                document.save(f"C:\\Users\\admin\Desktop\\Tallentelligent\\Practices_{index-12}.docx")

                # Used to end program once all the pages have been accounted for in files
                if page_num >= 390:
                    break
            else:
                # drivers portion of the pdf
                document.save(f"C:\\Users\\admin\Desktop\\Tallentelligent\\Driver_Marker_{index}.docx")
            index += 1

            # creates new document
            document = Document()


# def create_word_documents(text, page_groups):
#     """Creates Word documents based on the specified page groupings."""
#     current_document_index = 1
#     start_index = 0
#     for end_index in page_groups:
#         document_text = text[start_index:end_index]
#
#         start_index = end_index
#         current_document_index += 1

# Get user input for PDF file path and page groupings
pdf_file_path = "C:\\Users\\admin\Desktop\\Tallentelligent\\KSAP Field Guide v8.pdf"
page_groups = [22,32,42,52,62,72,82,92,102,112,124,134,144,154,164,174,184,194,204,216,226,236,246,256,266,276,286,296,306,316,326,336,346,356,366,376,386,397]

# page_groups = [397]

# Extract text from PDF
text = extract_text_from_pdf(pdf_file_path, page_groups)

# Create Word documents
# create_word_documents(text, page_groups)

print("Word documents created successfully!")
